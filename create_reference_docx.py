#!/usr/bin/env python3
"""
Create a reference DOCX file with Aptos font for pandoc.
This reference document will be used by pandoc to style the output.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_custom_style(styles, style_name, style_type, font_name='Calibri', font_size=11, 
                     is_superscript=False, base_style=None):
    """Add a custom style to the document if it doesn't exist."""
    try:
        # Try to get existing style
        style = styles[style_name]
        style.font.name = font_name
        if font_size:
            style.font.size = Pt(font_size)
        if is_superscript:
            style.font.superscript = True
        return style
    except KeyError:
        # Style doesn't exist, create it
        try:
            style = styles.add_style(style_name, style_type)
            style.font.name = font_name
            if font_size:
                style.font.size = Pt(font_size)
            if is_superscript:
                style.font.superscript = True
            if base_style:
                style.base_style = styles[base_style]
            return style
        except Exception as e:
            print(f"  Could not create style '{style_name}': {e}")
            return None

def create_reference_docx(output_file='reference_calibri.docx'):
    """Create a reference DOCX with Calibri font styling and justified paragraphs."""
    
    doc = Document()
    
    # Get styles
    styles = doc.styles
    
    # Set Normal style to use Calibri with justified alignment
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    # Set paragraph alignment to justified
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Set Title style
    try:
        title_style = styles['Title']
        title_font = title_style.font
        title_font.name = 'Calibri'
        title_font.size = Pt(24)
        title_font.bold = True
    except KeyError:
        pass
    
    # Set Heading 1 style
    try:
        heading1_style = styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = 'Calibri'
        heading1_font.size = Pt(16)
        heading1_font.bold = True
    except KeyError:
        pass
    
    # Set Heading 2 style
    try:
        heading2_style = styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.name = 'Calibri'
        heading2_font.size = Pt(14)
        heading2_font.bold = True
    except KeyError:
        pass
    
    # Set Heading 3 style
    try:
        heading3_style = styles['Heading 3']
        heading3_font = heading3_style.font
        heading3_font.name = 'Calibri'
        heading3_font.size = Pt(12)
        heading3_font.bold = True
    except KeyError:
        pass
    
    # Add/Set Footnote Reference style (the superscript number in text)
    footnote_ref = add_custom_style(styles, 'Footnote Reference', WD_STYLE_TYPE.CHARACTER, 
                                     font_name='Calibri', font_size=11, is_superscript=True)
    if footnote_ref:
        print("  Configured Footnote Reference style (superscript)")
    
    # Add/Set Footnote Text style (the actual footnote content)
    footnote_text = add_custom_style(styles, 'Footnote Text', WD_STYLE_TYPE.PARAGRAPH,
                                     font_name='Calibri', font_size=10)
    if footnote_text:
        print("  Configured Footnote Text style (10pt)")
    
    # Set Hyperlink style (for clickable links)
    try:
        hyperlink_style = styles['Hyperlink']
        hyperlink_font = hyperlink_style.font
        hyperlink_font.name = 'Calibri'
        hyperlink_font.color.rgb = None  # Use default blue
        hyperlink_font.underline = True
        print("  Configured Hyperlink style (blue, underlined)")
    except KeyError:
        # Create hyperlink style if it doesn't exist
        try:
            hyperlink_style = styles.add_style('Hyperlink', WD_STYLE_TYPE.CHARACTER)
            hyperlink_font = hyperlink_style.font
            hyperlink_font.name = 'Calibri'
            hyperlink_font.color.rgb = RGBColor(0, 0, 255)  # Blue
            hyperlink_font.underline = True
            print("  Created Hyperlink style (blue, underlined)")
        except Exception as e:
            print(f"  Warning: Could not configure Hyperlink style: {e}")
    
    # Set Body Text style
    try:
        body_text_style = styles['Body Text']
        body_text_font = body_text_style.font
        body_text_font.name = 'Calibri'
        body_text_font.size = Pt(11)
        # Set justified alignment for body text too
        body_text_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except KeyError:
        pass
    
    # Add a sample paragraph and footnote to ensure the document is valid
    para = doc.add_paragraph('This is a reference document for pandoc styling.')
    
    # Add a footnote to trigger footnote styles
    try:
        # Add sample text with a footnote
        sample_para = doc.add_paragraph('Sample text with footnote.')
        # Note: python-docx doesn't have direct footnote support, but the styles are still used
    except Exception as e:
        pass
    
    # Save the reference document
    doc.save(output_file)
    print(f"Created reference document: {output_file}")
    print(f"  Default font: Calibri")
    print(f"  Font size: 11pt")
    print(f"  Paragraph alignment: Justified")

if __name__ == '__main__':
    try:
        create_reference_docx()
    except ImportError:
        print("Error: python-docx package not found.")
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call(['pip3', 'install', 'python-docx'])
        print("Package installed. Please run the script again.")

